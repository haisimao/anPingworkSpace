# -*- coding: utf-8 -*-
# 只管解析,
import re
import zipfile
from io import StringIO
from io import BytesIO

import docx
import win32com
import xlrd
from bson import ObjectId
from win32com.client import Dispatch
import os, sys

# from tools.req_for_api import req_for_serial_number

curPath = os.path.abspath(os.path.dirname(__file__))
sys.path.append(curPath[:-9])
src_dir = curPath + "/js/"

import requests
import random
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter, process_pdf
from pdfminer.converter import PDFPageAggregator, TextConverter
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from PyPDF2 import PdfFileReader, PdfFileWriter
from urllib.request import urlopen
from SpiderTools.Tool import platform_system


FILE_PATH = ''


filesList = '''反复读这段代码并没有发现什么问题因为有些网页的附件名称是相同的例如所以我按每个网页的标题在总览页面爬到的分文件夹放置下载的文件所以方法中传了一个参数而如果参数传空则不会报错其实由此已经可以发现所在了但我却没想到又反复折腾了很久才发现原来是文件名太长了在下面单个文件名的长度限制是完整的路径长度如限制是路径最后有一个字符串结束符要占掉一个字符所以完整路径实际限长是'''


def parse_pdf(url=None):

    fp = urlopen(url)
    parser = PDFParser(fp)  # 用文件对象创建一个PDF文档分析器
    doc = PDFDocument() # 创建一个PDF文档
    parser.set_document(doc)  # 连接分析器，与文档对象
    doc.set_parser(parser)
    doc.initialize()  # 提供初始化密码，如果没有密码，就创建一个空的字符串
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建PDF，资源管理器，来共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # 循环遍历列表，每次处理一个page内容
        # doc.get_pages() 获取page列表
        result_list = list()
        for page in doc.get_pages():
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
            # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
            # 想要获取文本就获得对象的text属性，
            text_list = list()
            for x in layout:
                if isinstance(x, LTTextBoxHorizontal):
                    # with open(r'2.txt', 'a', encoding="utf-8") as f:
                    results = x.get_text()
                    text_list.append(results)
            text = "\n".join(text_list)
            result_list.append(text)

        return "".join(result_list)


def parse_zip(url=None):
    '''
    输入 .zip 文件, 解压文件, 返回被压缩文件的内容
    :param url:  Either the path to the file, or a file-like object.If it is a path, the file will be opened and closed by ZipFile.
    :return:
    '''
    if url:
        response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1', })
        url = BytesIO(response.content)

    zFile = zipfile.ZipFile(url)
    # ZipFile.namelist(): 获取ZIP文档内所有文件的名称列表
    str_data = ''
    for file_name in zFile.namelist():
        data = zFile.read(file_name)
        try:
            str_data += data.decode('gbk')
        except:
            try:
                str_data += str(data)
            except:
                str_data += ''
    zFile.close()
    return str_data


def parse_excel(url=None):
    '''

    :param url:  http://fgw.sz.gov.cn/fzggzl/zdxm/201902/P020190220602369247623.xls
    :return:
    '''
    # 由url获取 文件 , 写入文件夹
    if '\.xlsx' in url:
        file_name = src_dir + ''.join(random.choices(filesList, k=5)) + '.xlsx'
    else:
        file_name = src_dir + ''.join(random.choices(filesList, k=5)) + '.xls'
    response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1', })
    with open(file_name, 'wb') as fp:
        fp.write(response.content)  # 文件大多是二进制或者字符串传输的,

    wb = xlrd.open_workbook(filename=file_name)  # 打开文件
    # wb = xlrd.open_workbook(file_contents=response.text)  # 打开文件
    sheet_names = wb.sheet_names()[0] if isinstance(wb.sheet_names(), list) else wb.sheet_names()
    excel_data = wb.sheet_by_name(sheet_names)
    row = excel_data.nrows  # 总行数
    text = []
    for i in range(1, row):
        rowdata = excel_data.row_values(i)  # i行的list
        text.append(' '.join([str(_) for _ in rowdata]))
    os.remove(file_name)  # 删除请求的文件
    return '\n'.join(text)


def transform_data(url=None, data=None):
    type_ = data.get('conten_type')
    if type:
        if '.pdf' in type_:
            return parse_pdf(url)
        elif '.doc' in type_ or '.docx' in type_:
            return parse_word(url)
        elif '.zip' in type_:
            return parse_zip(url)
        elif '.xls' in type_ or '.xlsx' in type_:
            return parse_excel(url)
        else:
            return ' '
    else:
        return ' '


def find_type(url=None):
    if url:
        if '.pdf' in url:
            return '.pdf'
        elif '.docx' in url:
            return '.docx'
        elif '.doc' in url :
            return '.doc'
        elif '.zip' in url:
            return '.zip'
        elif '.xls' in url :
            return '.xls'
        elif '.xlsx' in url:
            return '.xlsx'
        else:
            return ''
    else:
        return ''


def doc2docx(url=None):
    '''
    :param url:  http://fgw.gz.gov.cn/gzplan/s15713/201902/ffff52619b4d4c40ae400f6ec873d1af/files/22ea53cb1e61499b9b6bb0162e5d8e65.doc
    :return: 新的文件名称
    '''
    # 由url获取 文件 , 写入文件夹
    try:
        file_name = src_dir + ''.join(random.choices(filesList, k=5)) + '.doc'
        response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1',}, timeout=200)
        with open(file_name, 'wb') as fp:
            fp.write(response.content)

        w = win32com.client.Dispatch('Word.Application')
        w.Visible = 0
        w.DisplayAlerts = 0

        doc = w.Documents.Open(file_name)  # 读取 .doc 文件
        newpath = os.path.splitext(file_name)[0] + '.docx'
        doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
        # doc.Close()
        w.Quit()
    except Exception as e:
        print(e)
        os.remove(file_name)
    else:
        # os.remove(file_name)  # 删除请求的文件, 若文件使用为关闭,将没有权限删除
        return newpath, file_name


def parse_word(url=None, type='doc'):
    '''

    :param url: 需要解析的 .doc 文件
    :return:
    '''
    if '.docx' not in url or type == 'doc':
        newpath, file_name = doc2docx(url)
        os.remove(file_name)  # 删除请求的文件
    else:
        response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1', })
        newpath = BytesIO(response.content)  # 数据读写不一定是文件，也可以在内存中读写. StringIO和BytesIO是在内存中操作str和bytes的方法

    doc = docx.Document(newpath)  # 可传入文件路径(字符串)或类似文件的对象(文件游标, 数据流)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text.strip() + '\n'
    # doc.Close()
    # doc.Quit()  # 好像不需要关闭, 官方文当中没有关闭

    os.remove(newpath) if isinstance(newpath, str) else ''  # 删除请求的文件
    return text


if __name__ == '__main__':
    # zip2txt('http://www.gz.gov.cn/publicfiles/business/htmlfiles/gzplanjg/cmsmedia/other/2013/5/other94483.zip')
    # print(parse_zip())
    # print(parse_excel('http://fgw.sz.gov.cn/fzggzl/zdxm/201902/P020190220602369247623.xls'))
    url = 'https://www.sz68.com/trade/download?module=base&service=BaseDownload&method=downloadFile&lobId=20200120092041483341001535481407'
    print(parse_word(url, 'doc'))
    # print(parse_pdf('http://drc.gd.gov.cn/attachment/0/112/112043/845093.pdf'))
    # 解决2003一起拿的 word 读取
    # file_name = src_dir + ''.join(random.choices(filesList, k=5)) + '.doc'
    # response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1', }, timeout=200)
    # with open(file_name, 'wb') as fp:
    #     print(response.content[300:-1200])
    #     fp.write(response.content[:-1200])
    # import glob
    # import subprocess
    #
    # for doc in glob.iglob("src_dir/*.doc"):
    #     subprocess.call(['soffice', '--headless', '--convert-to', 'docx', doc])


